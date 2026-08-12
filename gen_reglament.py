# -*- coding: utf-8 -*-
"""ROYD platformasi uchun Tashkiliy-Huquqiy Asos va Reglament hujjatini yaratish."""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

doc = Document()

# Sahifa sozlamalari
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)


def set_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper())
    set_font(run, size=14, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    set_font(run, size=12)
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    return p


def bullet(text, num=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    prefix = "%d." % num if num else "-"
    run = p.add_run("%s  %s" % (prefix, text))
    set_font(run, size=12)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_bold_body(label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    r1 = p.add_run(label)
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=12)
    p.paragraph_format.space_after = Pt(3)
    return p


# ==============================================================================
# MUQOVA SAHIFASI
# ==============================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM, FAN VA INNOVATSIYALAR VAZIRLIGI")
set_font(r, size=12, bold=True)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("[UNIVERSITET NOMI]")
set_font(r, size=13, bold=True)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("TASDIQLANDI")
set_font(r, size=12, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Rektor ___________________ [F.I.O.]")
set_font(r, size=12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("<<____>> _____________ 2025-yil")
set_font(r, size=12)

for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "REGISTRATOR OFIS - YAGONA DARCHA TIZIMI (ROYD)\n"
    "TASHKILIY-HUQUQIY ASOSI VA FOYDALANISH REGLAMENTI"
)
set_font(r, size=16, bold=True)

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("2025-yil")
set_font(r, size=12)

doc.add_page_break()

# ==============================================================================
# 1-BOB. UMUMIY QOIDALAR
# ==============================================================================
heading1("1-BOB. UMUMIY QOIDALAR")

heading2("1.1. Hujjatning maqsadi")
body(
    "Ushbu reglament <<Registrator Ofis - Yagona Darcha Tizimi>> (keyingi o'rinlarda - ROYD yoki "
    "Tizim) nomli elektron platformaning tashkiliy-huquqiy asosini belgilab, barcha "
    "foydalanuvchilar - talabalar, registratorlar, xodimlar va rahbariyat uchun tizimdan "
    "foydalanish qoidalarini tartibga soladi."
)

heading2("1.2. Qo'llanish doirasi")
body(
    "Ushbu reglament universitetning barcha tarkibiy bo'limlari, fakultetlar, kafedra xodimlari, "
    "o'quv ishlari bo'yicha bo'lim xodimlari, registratorlar hamda universitetda tahsil olayotgan "
    "barcha talabalar va magistrantlarga nisbatan qo'llaniladi."
)

heading2("1.3. Asosiy tushunchalar")

terms = [
    ("ROYD: ", "Yagona Darcha Tizimi - talabalar murojaatlarini elektron shaklda qabul qilish, yo'naltirish va kuzatish platformasi."),
    ("Murojaat: ", "Talaba tomonidan Tizim orqali yuboriladigan so'rov, ariza yoki shikoyat."),
    ("Tracking-raqam: ", "Har bir murojaatga avtomatik beriluvchi noyob identifikatsiya raqami (masalan: ROYD-2025-000123)."),
    ("SLA: ", "Service Level Agreement - murojaatni ko'rib chiqish uchun belgilangan maksimal muddat."),
    ("KPI: ", "Key Performance Indicators - registrator va xodimlar samaradorligini o'lchash ko'rsatkichlari."),
    ("HEMIS: ", "O'zbekiston OTM talabalari ma'lumotlar tizimi; talabalar ushbu tizim orqali autentifikatsiya qilinadi."),
    ("Registrator: ", "Murojaatlarni qabul qilib, tegishli bo'limlarga yo'naltiruvchi mas'ul xodim."),
    ("Status: ", "Murojaatning joriy holati: Yangi / Qabul qilindi / Ko'rib chiqilmoqda / Bajarildi / Rad etildi / Qaytarildi."),
]
for label, definition in terms:
    add_bold_body(label, definition)

doc.add_page_break()

# ==============================================================================
# 2-BOB. TASHKILIY-HUQUQIY ASOS
# ==============================================================================
heading1("2-BOB. TASHKILIY-HUQUQIY ASOS")

heading2("2.1. Qonunchilik asoslari")
body("ROYD tizimi quyidagi me'yoriy-huquqiy hujjatlar asosida ishlab chiqilgan va joriy etilgan:")

legal_bases = [
    "O'zbekiston Respublikasi Konstitutsiyasi (2023-yil yangi tahriri), 43-modda - har kim o'z huquq va erkinliklarini qonun bilan taqiqlanmagan barcha usullar bilan himoya qilishga haqli.",
    "\"Axborot erkinligi prinsiplari va ularni ta'minlashning kafolatlari to'g'risida\" O'zbekiston Respublikasi Qonuni (2002-yil 12-dekabr, 439-II-son).",
    "\"Elektron hukumat to'g'risida\" O'zbekiston Respublikasi Qonuni (2015-yil 28-aprel, 389-V-son).",
    "\"Axborot texnologiyalari va axborotni himoya qilish sohasidagi faoliyatni tartibga solish to'g'risida\" O'zbekiston Respublikasi Qonuni.",
    "O'zbekiston Respublikasi Prezidentining \"Oliy ta'lim tizimini yanada takomillashtirish chora-tadbirlari to'g'risida\" PQ-5847-son Qarori (2019-yil 8-oktyabr).",
    "Oliy ta'lim, fan va innovatsiyalar vazirligi tomonidan tasdiqlangan \"Oliy ta'lim muassasalarida elektron hujjat aylanmasini tashkil etish tartibi to'g'risida\" Yo'riqnoma.",
    "\"Shaxsiy ma'lumotlar to'g'risida\" O'zbekiston Respublikasi Qonuni (2019-yil 2-iyul, 554-XIV-son) - talabalar shaxsiy ma'lumotlarini himoya qilish majburiyati.",
    "Universitetning Ustavi va ichki tartib qoidalari.",
]
for i, item in enumerate(legal_bases, 1):
    bullet(item, num=i)

heading2("2.2. Ichki me'yoriy hujjatlar")
body("Ushbu reglamentdan tashqari quyidagi hujjatlar Tizimning huquqiy asosini to'ldiradi:")
internal_docs = [
    "Universitetda elektron xizmatlarni joriy etish bo'yicha Rektor buyrug'i.",
    "Registratorlar bo'limi to'g'risidagi Nizom.",
    "Talabalar bilan ishlash bo'yicha Metodologik ko'rsatma.",
    "Axborot xavfsizligi siyosati.",
    "Shaxsiy ma'lumotlarga ishlov berish siyosati (Privacy Policy).",
]
for i, doc_item in enumerate(internal_docs, 1):
    bullet(doc_item, num=i)

heading2("2.3. Tizimni joriy etish")
body(
    "ROYD tizimi rektor buyrug'i asosida rasman joriy etiladi. Tizimga o'zgartirishlar kiritish, "
    "muddatlarni qayta ko'rib chiqish yoki yangi funksionallik qo'shish rektor yoki uning "
    "o'rinbosarining roziligi bilan amalga oshiriladi."
)

doc.add_page_break()

# ==============================================================================
# 3-BOB. TIZIMNING MAQSAD VA VAZIFALARI
# ==============================================================================
heading1("3-BOB. TIZIMNING MAQSAD VA VAZIFALARI")

heading2("3.1. Asosiy maqsad")
body(
    "ROYD platformasining asosiy maqsadi - universitetda talabalarning ariza va murojaatlarini "
    "qog'ozsiz, shaffof va tezkor tarzda ko'rib chiqishni ta'minlash, xizmat ko'rsatish "
    "sifatini oshirish va ma'muriy yuklarni kamaytirish."
)

heading2("3.2. Asosiy vazifalar")
tasks = [
    "Talabalar murojaatlarini yagona elektron kanal orqali qabul qilish va ro'yxatga olish.",
    "Har bir murojaatga noyob tracking-raqam berish va uning holatini real vaqt rejimida kuzatish imkonini ta'minlash.",
    "Murojaatlarni kategoriyalar bo'yicha avtomatik saralash va mas'ul bo'limlarga yo'naltirish.",
    "Belgilangan SLA muddatlarini nazorat qilish va muddatlarning o'tkazib yuborilishini oldini olish.",
    "Registratorlar va xodimlar samaradorligini KPI ko'rsatkichlari orqali baholash.",
    "Talaba va xodim o'rtasida xavfsiz yozishmalar olib borish imkonini yaratish.",
    "Rahbariyatga analitik hisobotlar va statistikani taqdim etish.",
    "HEMIS tizimi bilan integratsiya orqali talabalar ma'lumotlarini avtomatik olish.",
    "Bildirishnomalar (elektron pochta, Telegram) orqali barcha tomonlarni xabardor qilish.",
    "Barcha amallar bo'yicha audit izini saqlash.",
]
for i, task in enumerate(tasks, 1):
    bullet(task, num=i)

doc.add_page_break()

# ==============================================================================
# 4-BOB. FOYDALANUVCHILAR VA ULARNING HUQUQLARI
# ==============================================================================
heading1("4-BOB. FOYDALANUVCHILAR VA ULARNING HUQUQLARI")

heading2("4.1. Foydalanuvchi rollari")
body(
    "Tizimda beshta asosiy rol mavjud. Har bir rol o'ziga xos vakolatlar va "
    "ma'lumotlarga kirish huquqlariga ega."
)

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["Rol", "Tizim nomi", "Asosiy vakolatlar"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    set_font(r, size=11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

roles_data = [
    ("Talaba", "student",
     "Murojaat yaratish; o'z murojaatlarini ko'rish; xabarnomalar olish; xodim bilan yozishish."),
    ("Registrator", "registrator",
     "Barcha murojaatlarni ko'rish; qabul qilish/rad etish/yo'naltirish; xodimga belgilash; SLA kuzatuvi."),
    ("Xodim (Staff)", "staff",
     "O'ziga biriktirilgan murojaatlarni ko'rish; holat yangilash; talaba bilan yozishish; fayl yuklash."),
    ("Rahbariyat", "leadership",
     "Barcha ma'lumotlarni o'qish; analitik hisobotlar; KPI paneli; statistika."),
    ("Administrator", "admin",
     "To'liq boshqaruv; foydalanuvchi va rol boshqaruvi; kategoriya va bo'lim sozlamalari; tizim konfiguratsiyasi."),
]
for role, sys_name, perms in roles_data:
    row = table.add_row().cells
    row[0].paragraphs[0].add_run(role)
    set_font(row[0].paragraphs[0].runs[0], size=11)
    row[1].paragraphs[0].add_run(sys_name)
    set_font(row[1].paragraphs[0].runs[0], size=11)
    row[2].paragraphs[0].add_run(perms)
    set_font(row[2].paragraphs[0].runs[0], size=10)

doc.add_paragraph()

heading2("4.2. Tizimga kirish")
body(
    "Talabalar HEMIS tizimidagi talaba ID raqami va parol orqali autentifikatsiya qilinadi. "
    "Tizim HEMIS API orqali talabaning ma'lumotlarini (ism, guruh, fakultet, mutaxassislik, "
    "kurs va boshqalar) avtomatik ravishda oladi."
)
body(
    "Registratorlar, xodimlar, rahbariyat va administratorlar elektron pochta manzili va parol "
    "orqali tizimga kiradi. Parollar bcrypt algoritmi orqali shifrlanadi va ochiq ko'rinishda "
    "saqlanmaydi."
)

heading2("4.3. Foydalanuvchilarning majburiyatlari")
body("Barcha foydalanuvchilar quyidagi majburiyatlarni bajarishlari shart:")
obligations = [
    "Login va parolni uchinchi shaxslarga bermaydi.",
    "Tizimdan faqat o'z vakolatlari doirasida foydalanadi.",
    "Yolg'on yoki asos bo'lmagan murojaatlar yuborilmaydi.",
    "Boshqa foydalanuvchilarning ma'lumotlariga ruxsatsiz kirishga urinmaydi.",
    "Tizimda muloqot madaniyatiga rioya qiladi; haqoratli yoki tahdid ifodalovchi xabarlar yuborilmaydi.",
    "Qonunchilikka zid kontentlar yuklanmaydi.",
]
for i, ob in enumerate(obligations, 1):
    bullet(ob, num=i)

doc.add_page_break()

# ==============================================================================
# 5-BOB. MUROJAAT YUBORISH TARTIBI
# ==============================================================================
heading1("5-BOB. MUROJAAT YUBORISH TARTIBI (TALABA UCHUN)")

heading2("5.1. Murojaat kategoriyalari")
body(
    "Tizimda murojaatlar quyidagi asosiy kategoriyalar bo'yicha guruhlanadi (administrator "
    "tomonidan boshqariladi va kengaytirilishi mumkin):"
)
categories = [
    ("Akademik masalalar: ", "Qayta o'qish, imtihon natijasi, reyting hisob-kitobi."),
    ("Ma'lumotnomalar: ", "O'qish ma'lumotnomasi, akademik spravka, dekanat ma'lumotnomalari."),
    ("Moliyaviy masalalar: ", "To'lov muddatlari, to'lov tizimi, stipendiya savollari."),
    ("Ijtimoiy va turar-joy: ", "Yotoqxona ariza, sotsial nafaqa, imtiyozlar."),
    ("Ruxsat va ta'til: ", "Akademik ta'til, arxiv, ishdanartiqlarga ruxsat."),
    ("Shikoyatlar: ", "Xodim yoki jarayon bo'yicha rasmiy shikoyat."),
    ("Boshqa murojaatlar: ", "Yuqoridagi kategoriyalarga kirmaydigan so'rovlar."),
]
for label, desc in categories:
    add_bold_body(label, desc)

heading2("5.2. Murojaat yaratish bosqichlari")
steps = [
    "Tizimga HEMIS ID va parol bilan kirish.",
    "\"Yangi murojaat\" tugmasini bosish.",
    "Kategoriya tanlash.",
    "Sarlavha (maksimal 500 belgi) va batafsil tavsif (majburiy) yozish.",
    "Zarur bo'lsa, qo'shimcha fayl(lar) biriktirish (ruxsat etilgan formatlar: PDF, DOCX, XLSX, JPG, PNG; maksimal hajm: 10 MB).",
    "Murojaatni yuborish.",
    "Tizim avtomatik ravishda tracking-raqam beradi va tasdiqlash xabari yuboriladi.",
]
for i, step in enumerate(steps, 1):
    bullet(step, num=i)

heading2("5.3. Murojaatni kuzatish")
body(
    "Talaba istalgan vaqt tizimga kirib, o'z murojaatining joriy holatini, mas'ul xodimni, "
    "tarixini va yozishmalarini ko'rishi mumkin. Status o'zgarganda elektron pochta orqali "
    "avtomatik bildirishnoma keladi."
)

doc.add_page_break()

# ==============================================================================
# 6-BOB. MUROJAATLARNI KO'RIB CHIQISH REGLAMENTI
# ==============================================================================
heading1("6-BOB. MUROJAATLARNI KO'RIB CHIQISH REGLAMENTI")

heading2("6.1. Ish jarayoni (Workflow)")
workflow = [
    ("YANGI (new): ", "Talaba murojaatni yubordi. Registratorga bildirishnoma boradi."),
    ("QABUL QILINDI (accepted): ", "Registrator murojaatni ko'rib, asosli ekanini tasdiqladi."),
    ("KO'RIB CHIQILMOQDA (in_progress): ", "Mas'ul xodimga biriktirildi, ish boshlangan."),
    ("QAYTARILDI (returned): ", "Qo'shimcha ma'lumot yoki hujjat talab qilinadi. Talaba to'ldirishi shart."),
    ("BAJARILDI (completed): ", "Murojaat to'liq hal etildi."),
    ("RAD ETILDI (rejected): ", "Murojaat asoslanmagan yoki sababli rad etildi."),
]
for label, desc in workflow:
    add_bold_body(label, desc)

heading2("6.2. Registratorning vazifalari")
reg_duties = [
    "Yangi murojaatlarni 1 ish kuni ichida ko'rib chiqish.",
    "Asoslanmagan yoki kategoriyasi noto'g'ri murojaatni sababini ko'rsatib qaytarish.",
    "Murojaatni tegishli xodim yoki bo'limga biriktirish.",
    "SLA muddatlarini nazorat qilish va kechikishlar haqida rahbariyatni xabardor qilish.",
    "Tizimda barcha amallar bo'yicha izoh qoldirish.",
]
for i, d in enumerate(reg_duties, 1):
    bullet(d, num=i)

heading2("6.3. Xodimning vazifalari")
staff_duties = [
    "Biriktirilgan murojaatni belgilangan SLA muddat ichida ko'rib chiqish.",
    "Zarur holda talabadan qo'shimcha ma'lumot so'rash (murojaat <<Qaytarildi>> statusiga o'tkaziladi).",
    "Natija va qaror haqida tizimda izoh qoldirish.",
    "Bajarilgan murojaatni <<Bajarildi>> statusiga o'tkazish.",
]
for i, d in enumerate(staff_duties, 1):
    bullet(d, num=i)

heading2("6.4. Prioritet darajalari")
priorities = [
    ("Oddiy (normal): ", "Standart SLA muddati qo'llaniladi."),
    ("Yuqori (high): ", "SLA muddati 50% qisqartiriladi; registratorga darhol bildirishnoma."),
    ("Favqulodda (urgent): ", "SLA muddati 75% qisqartiriladi; registrator va bo'lim boshlig'iga bildirishnoma."),
]
for label, desc in priorities:
    add_bold_body(label, desc)

doc.add_page_break()

# ==============================================================================
# 7-BOB. SLA
# ==============================================================================
heading1("7-BOB. SLA - XIZMAT KO'RSATISH DARAJASI SHARTNOMASI")

heading2("7.1. SLA muddatlari jadvali")
body(
    "Quyida har bir murojaat kategoriyasi uchun belgilangan standart SLA muddatlari "
    "ko'rsatilgan (ish kunlari hisobida):"
)

sla_table = doc.add_table(rows=1, cols=4)
sla_table.style = "Table Grid"
sla_table.alignment = WD_TABLE_ALIGNMENT.CENTER
sla_hdr = sla_table.rows[0].cells
for cell, txt in zip(sla_hdr, ["Kategoriya", "Oddiy (kun)", "Yuqori (kun)", "Favqulodda (kun)"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    set_font(r, size=11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sla_data = [
    ("Akademik masalalar",  "5", "3", "2"),
    ("Ma'lumotnomalar",     "3", "2", "1"),
    ("Moliyaviy masalalar", "5", "3", "2"),
    ("Ijtimoiy va turar-joy","7","4", "2"),
    ("Ruxsat va ta'til",    "5", "3", "2"),
    ("Shikoyatlar",         "10","5", "3"),
    ("Boshqa murojaatlar",  "5", "3", "2"),
]
for cat, n, h, u in sla_data:
    row = sla_table.add_row().cells
    for cell, val in zip(row, [cat, n, h, u]):
        p = cell.paragraphs[0]
        r = p.add_run(val)
        set_font(r, size=11)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

heading2("7.2. SLA buzilishi tartibi")
body(
    "SLA muddati tugashiga 24 soat qolganda xodimga eslatma yuboriladi. "
    "Muddat o'tkazib yuborilganda registrator va bo'lim boshlig'iga avtomatik bildirishnoma boradi. "
    "Bunday holatlar KPI hisobotiga salbiy ko'rsatkich sifatida kiritiladi."
)

heading2("7.3. Dam olish kunlari")
body(
    "SLA muddatlarini hisoblashda faqat ish kunlari (dushanba-juma) hisobga olinadi. "
    "Rasmiy bayram kunlari ham ish kuni hisoblanmaydi. "
    "SLA muddat o'rnatish algoritmida ushbu qoidalar avtomatik qo'llaniladi."
)

doc.add_page_break()

# ==============================================================================
# 8-BOB. KPI
# ==============================================================================
heading1("8-BOB. KPI - SAMARADORLIK KO'RSATKICHLARI")

heading2("8.1. Registratorlar uchun KPI")
kpi_reg = [
    ("Qabul qilish tezligi: ", "Yangi murojaat o'rtacha qabul qilish vaqti (maqsad: 1 ish kuni)."),
    ("SLA saqlanishi: ", "O'z vaqtida ko'rib chiqilgan murojaatlar foizi (maqsad: 95% va yuqori)."),
    ("Rad etish darajasi: ", "Asossiz rad etilgan murojaatlar foizi (maqsad: 5% dan kam)."),
    ("Qayta ishlov ko'rsatkichi: ", "Bir marta <<Qaytarildi>>ga o'tgan murojaatlar foizi."),
]
for label, desc in kpi_reg:
    add_bold_body(label, desc)

heading2("8.2. Xodimlar uchun KPI")
kpi_staff = [
    ("Hal qilish tezligi: ", "O'rtacha murojaatni yopish vaqti (SLA ga nisbatan foizda)."),
    ("SLA saqlanishi: ", "Belgilangan muddat ichida yopilgan murojaatlar foizi (maqsad: 90% va yuqori)."),
    ("Talaba qoniqish darajasi: ", "Yopilgan murojaatlar bo'yicha talaba baholash reytingi (kelajak bosqich)."),
    ("Qayta ochilgan murojaatlar: ", "Yopilgandan so'ng yana ochilgan murojaatlar foizi."),
]
for label, desc in kpi_staff:
    add_bold_body(label, desc)

heading2("8.3. KPI hisobotlari")
body(
    "Rahbariyat va administrator oylik KPI hisobotini tizimdan yuklab olishi mumkin. "
    "Hisobotlar fakultet, bo'lim, registrator va xodim kesimida taqdim etiladi. "
    "KPI natijalari xodimlarning mehnat samaradorligini baholashda qo'shimcha ma'lumot "
    "sifatida ishlatilishi mumkin."
)

doc.add_page_break()

# ==============================================================================
# 9-BOB. FAYL VA HUJJAT BOSHQARUVI
# ==============================================================================
heading1("9-BOB. FAYL VA HUJJAT BOSHQARUVI")

heading2("9.1. Ruxsat etilgan fayl formatlari")
body("Tizimga quyidagi formatlardagi fayllar yuklanishi mumkin:")
formats = [
    "Hujjatlar: PDF, DOCX, DOC, XLSX, XLS, PPTX",
    "Rasmlar: JPG, JPEG, PNG, GIF",
    "Maksimal fayl hajmi: 10 MB (bitta fayl)",
    "Bitta murojaatda maksimal fayl soni: 5 ta",
]
for f in formats:
    bullet(f)

heading2("9.2. Fayl saqlash")
body(
    "Yuklangan fayllar server tomonida xavfsiz saqlash katalogida saqlanadi. "
    "Fayl yo'llari ma'lumotlar bazasida saqlanadi. "
    "Fayllar faqat tegishli murojaat ishtirokchilari (talaba, biriktirilgan xodim, registrator) "
    "tomonidan yuklab olinishi mumkin."
)

heading2("9.3. Saqlash muddati")
body(
    "Yopilgan murojaatlar va ularga biriktirilgan fayllar arxivda kamida 5 yil saqlanadi. "
    "O'chirilgan murojaatlar audit jurnaliga yozib qo'yiladi."
)

# ==============================================================================
# 10-BOB. MA'LUMOT XAVFSIZLIGI VA MAXFIYLIK
# ==============================================================================
heading1("10-BOB. MA'LUMOT XAVFSIZLIGI VA MAXFIYLIK")

heading2("10.1. Texnik xavfsizlik choralari")
security_measures = [
    "Barcha ulanishlar HTTPS (TLS 1.2+) orqali amalga oshiriladi.",
    "Parollar bcrypt algoritmi (salt rounds 12 va yuqori) bilan xeshlangan holda saqlanadi.",
    "JWT (JSON Web Token) texnologiyasi orqali sessiyalar boshqariladi; access token muddati 30 daqiqa.",
    "Redis keshida sessiya tokenlari saqlanadi va muddati tugaganda avtomatik o'chiriladi.",
    "Har bir API so'rov autentifikatsiya va avtorizatsiya tekshiruvidan o'tadi.",
    "SQL injection, XSS va CSRF hujumlariga qarshi himoya sozlangan.",
    "Ma'lumotlar bazasiga to'g'ridan-to'g'ri kirish faqat ichki tarmoq orqali mumkin.",
    "Infratuzilma Docker konteynerlarida izolyatsiya qilingan holda ishlaytiladi.",
]
for i, m in enumerate(security_measures, 1):
    bullet(m, num=i)

heading2("10.2. Shaxsiy ma'lumotlarni himoya qilish")
body(
    "ROYD tizimi <<Shaxsiy ma'lumotlar to'g'risida>> O'zbekiston Respublikasi Qonuni talablariga "
    "muvofiq ishlaydi. Talabalar shaxsiy ma'lumotlari (ism, tug'ilgan sana, manzil, rasm va h.k.) "
    "faqat xizmat ko'rsatish maqsadida ishlatiladi va uchinchi shaxslarga berilmaydi."
)

doc.add_page_break()

# ==============================================================================
# 11-BOB. BILDIRISHNOMALAR TARTIBI
# ==============================================================================
heading1("11-BOB. BILDIRISHNOMALAR TARTIBI")

heading2("11.1. Bildirishnoma kanallari")
body("Tizim quyidagi kanallar orqali bildirishnoma yuboradi:")
notif_channels = [
    "Elektron pochta (e-mail) - asosiy kanal; barcha murojaatlar bo'yicha.",
    "Telegram bot - talabalar va xodimlar ulashlari mumkin bo'lgan qo'shimcha kanal (2-bosqich).",
    "Tizim ichidagi bildirishnomalar paneli - tizimga kirgan foydalanuvchi uchun.",
]
for n in notif_channels:
    bullet(n)

heading2("11.2. Bildirishnoma hodisalari")
notif_events = [
    ("Murojaat yaratildi: ", "Talabaga tracking-raqam va tasdiqlash yuboriladi."),
    ("Murojaat qabul qilindi: ", "Talabaga xabardorlik, registratorga yangi murojaat signali."),
    ("Murojaat biriktirildi: ", "Xodimga biriktirilgan murojaat haqida bildirishnoma."),
    ("Status o'zgarishi: ", "Talabaga har bir status o'zgarishi haqida xabar."),
    ("Murojaat qaytarildi: ", "Talabaga izoh bilan qaytarilganligi haqida xabar."),
    ("Yangi xabar: ", "Yozishmada yangi xabar kelganda qarshi tomonga bildirishnoma."),
    ("SLA eslatmasi: ", "Muddatga 24 soat qolganda mas'ul xodimga eslatma."),
    ("SLA o'tkazib yuborildi: ", "Registrator va bo'lim boshlig'iga tezkor bildirishnoma."),
    ("Murojaat yopildi: ", "Talabaga yakuniy natija va izoh bilan bildirishnoma."),
]
for label, desc in notif_events:
    add_bold_body(label, desc)

# ==============================================================================
# 12-BOB. AUDIT VA MONITORING
# ==============================================================================
heading1("12-BOB. AUDIT VA MONITORING")

heading2("12.1. Audit jurnali")
body(
    "Tizimda barcha muhim amallar audit jurnaliga yozib boriladi. Audit yozuvlari quyidagi "
    "ma'lumotlarni o'z ichiga oladi: amalga oshirgan foydalanuvchi, amal turi, vaqt belgisi, "
    "IP-manzil, o'zgartirilgan ma'lumot. Audit jurnali o'chirib bo'lmaydi."
)

heading2("12.2. Monitoring")
body(
    "Administrator panelida real vaqtda quyidagi ko'rsatkichlar kuzatiladi: "
    "faol foydalanuvchilar soni, yangi/ochiq/yopilgan murojaatlar soni, "
    "o'rtacha ko'rib chiqish vaqti, SLA bajarilish foizi va tizim holati (uptime)."
)

doc.add_page_break()

# ==============================================================================
# 13-BOB. JAVOBGARLIK VA NIZOLARNI HAL ETISH
# ==============================================================================
heading1("13-BOB. JAVOBGARLIK VA NIZOLARNI HAL ETISH")

heading2("13.1. Ma'muriy javobgarlik")
body(
    "Ushbu reglamentni buzgan xodimlar O'zbekiston Respublikasi mehnat qonunchiligiga muvofiq "
    "intizomiy javobgarlikka tortilishi mumkin. Shaxsiy ma'lumotlarni ruxsatsiz oshkor qilish "
    "<<Shaxsiy ma'lumotlar to'g'risida>> qonun doirasida huquqiy javobgarlikni keltirib chiqaradi."
)

heading2("13.2. Nizolarni hal etish tartibi")
dispute_steps = [
    "Talaba yoki xodim birinchi navbatda tizim orqali registratorga murojaat qiladi.",
    "Registrator 2 ish kuni ichida nizoni ko'rib chiqadi.",
    "Nizo hal bo'lmasa, masala bo'lim boshlig'i yoki dekan darajasiga ko'tariladi.",
    "Keyingi bosqich - prorektoral murojaat.",
    "Qonuniy da'vo holatlarida O'zbekiston Respublikasi qonunchiligiga muvofiq sud tartibida ko'rib chiqiladi.",
]
for i, step in enumerate(dispute_steps, 1):
    bullet(step, num=i)

# ==============================================================================
# 14-BOB. YAKUNIY QOIDALAR
# ==============================================================================
heading1("14-BOB. YAKUNIY QOIDALAR")

body(
    "Ushbu reglament rektor tomonidan tasdiqlangan kundan e'tiboran kuchga kiradi va "
    "universitetning ichki me'yoriy hujjati hisoblanadi."
)
body(
    "Reglamentga o'zgartirish va qo'shimchalar kiritish rektor buyrug'i asosida amalga oshiriladi "
    "va Tizim administratori tomonidan amalga tatbiq etiladi."
)
body(
    "Ushbu reglament talabalar va xodimlar uchun onlayn (Tizim ichida) hamda qog'oz shaklida "
    "mavjud bo'ladi. Yangi xodimlar ish boshlashdan oldin reglament bilan tanishadilar va "
    "imzoli tasdiqlash varag'iga qo'l qo'yadilar."
)
body(
    "Ushbu reglament Tizim joriy etilgunga qadar mavjud bo'lgan qog'oz shaklida murojaatlarni "
    "ko'rib chiqish tartiblarini bekor qiladi."
)

doc.add_page_break()

# ==============================================================================
# ILOVALAR
# ==============================================================================
heading1("ILOVALAR")

heading2("1-ilova. Murojaat holatlari diagrammasi")

statuses = [
    ("YANGI", "QABUL QILINDI  yoki  RAD ETILDI"),
    ("QABUL QILINDI", "KO'RIB CHIQILMOQDA"),
    ("KO'RIB CHIQILMOQDA", "BAJARILDI  yoki  RAD ETILDI  yoki  QAYTARILDI"),
    ("QAYTARILDI", "KO'RIB CHIQILMOQDA  (talaba ma'lumot qo'shganidan so'ng)"),
]
for from_s, to_s in statuses:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    r1 = p.add_run("%s  -->  " % from_s)
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(to_s)
    set_font(r2, size=12)
    p.paragraph_format.space_after = Pt(3)

heading2("2-ilova. Standart xabar shablonlari")

msg_templates = [
    (
        "Murojaat yaratildi: ",
        "Hurmatli [Ism], murojaatingiz qabul qilindi. Tracking-raqam: [ROYD-XXXX-XXXXXX]. "
        "Murojaatingizni tizim orqali kuzatishingiz mumkin."
    ),
    (
        "Murojaat yopildi: ",
        "Hurmatli [Ism], [ROYD-XXXX-XXXXXX] tracking-raqamli murojaatingiz ko'rib chiqildi. "
        "Natija: [Natija tavsifi]. Savollar bo'lsa, tizim orqali murojaat qilishingiz mumkin."
    ),
    (
        "Murojaat qaytarildi: ",
        "Hurmatli [Ism], [ROYD-XXXX-XXXXXX] tracking-raqamli murojaatingiz qo'shimcha ma'lumot "
        "talab etilganligi sababli qaytarildi. Sabab: [Sabab]. Iltimos, tizimga kirib "
        "so'ralgan ma'lumotni qo'shing."
    ),
]
for label, text in msg_templates:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.25)
    r1 = p.add_run(label)
    set_font(r1, size=12, bold=True)
    r2 = p.add_run(text)
    set_font(r2, size=12, italic=True)
    p.paragraph_format.space_after = Pt(4)

heading2("3-ilova. Reglamentni tasdiqlash varag'i")
body("Ushbu reglament bilan tanishib, uni bajarishga rozilik bildiraman:", indent=False)
doc.add_paragraph()

sign_table = doc.add_table(rows=1, cols=3)
sign_table.style = "Table Grid"
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for cell, txt in zip(sign_table.rows[0].cells, ["F.I.O.", "Lavozim", "Imzo va sana"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    set_font(r, size=11, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

for _ in range(6):
    row = sign_table.add_row().cells
    for cell in row:
        cell.paragraphs[0].add_run("                                       ")

# ------------------------------------------------------------------------------
output_path = "/home/ahror/Projects/registrator_ofis/ROYD_Tashkiliy_Huquqiy_Asos_va_Reglament.docx"
doc.save(output_path)
print("Hujjat saqlandi: %s" % output_path)
